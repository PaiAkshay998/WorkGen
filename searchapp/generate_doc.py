from docx import Document


def convert_to_doc(dict):
    document = Document()
    document.add_heading("Questions sheet")
    p=document.add_paragraph()
    p.add_run("set 1 marks=1").bold=True
    i=0
    for row in dict:
        if row.get('question_type')=='1A':
            p=document.add_paragraph()
            p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
            for question in row['question']:
                i=i+1
                p1=document.add_paragraph(str(i)+": ")
                p1.add_run(question)

    p=document.add_paragraph()
    p.add_run("set 2 marks=1").bold=True

    for row in dict:
        if row['question_type']=='1B':
            p=document.add_paragraph()
            p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
            for question in row['question']:
                i=i+1
                p1=document.add_paragraph(str(i)+": ")
                p1.add_run(question)

    p=document.add_paragraph()
    p.add_run("set 3 marks=2").bold=True

    for row in dict:
        if row['question_type']=='2':
            p=document.add_paragraph()
            p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
            for question in row['question']:
                i=i+1
                p1=document.add_paragraph(str(i)+": ")
                p1.add_run(question)

    p=document.add_paragraph()
    p.add_run("set 4 marks=3").bold = True
    for row in dict:
        if row['question_type']=='3':
            p=document.add_paragraph()
            p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
            for question in row['question']:
                i=i+1
                p1=document.add_paragraph(str(i)+" : ")
                p1.add_run(question)

    p=document.add_paragraph()
    p.add_run("set 5 marks=4").bold=True

    for row in dict:
        if row['question_type']=='4':
            p=document.add_paragraph()
            p.add_run("Attempt only "+str(row['attempt'])+" questions").italic=True
            for question in row['question']:
                i=i+1
                p1=document.add_paragraph(str(i)+" : ")
                p1.add_run(question)


    document.save("test_worsheet.docx")